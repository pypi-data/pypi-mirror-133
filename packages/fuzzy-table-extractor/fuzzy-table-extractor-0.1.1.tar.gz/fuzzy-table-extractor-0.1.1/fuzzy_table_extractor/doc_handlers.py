import logging
import os
import re
from abc import ABC, abstractmethod
from functools import lru_cache
from pathlib import Path
from typing import List

import numpy as np
import pandas as pd
import win32com.client as win32
from docx.api import Document
from docx.table import Table

# TODO move this to config file
AUX_FOLDER_PATH = r"C:\Users\leona\Documents\Repositories\FuzzyTableExtractor\temp"


class BaseHandler(ABC):
    def __init__(self, file_path: Path) -> None:
        self.file_path = file_path

    @property
    def words(self) -> List[str]:
        """List of all words in document"""
        return []

    @property
    def tables(self) -> List[pd.DataFrame]:
        """List of all tables (as dataframes) in document"""
        return []

    @property
    def dictionary(self) -> pd.DataFrame:
        """All cell couples in document"""
        return pd.DataFrame()


class DocxHandler(BaseHandler):
    def __init__(self, file_path: Path):
        super().__init__(file_path)
        self.__file_path = file_path

    @property
    @lru_cache
    def words(self) -> List[str]:
        logging.log(logging.INFO, "Getting words from document")
        document = self.document
        words = []

        # TODO check if there is another source of text that can be used for data extraction

        # Adding text from paragraphs
        paragraphs = document.paragraphs
        for item in paragraphs:
            aux = item.text
            aux = aux.split()
            words.extend(aux)

        # Adding text from tables
        for table in document.tables:
            for row in table.rows:
                data = [cell.text for cell in row.cells]
                for item in data:
                    words.extend(item.split())

        words = list(set(words))
        words = [word.lower() for word in words]
        words = [re.sub(r"[,]$|[.]$|[\(|\)|\:|\;|\'|\"]", "", word) for word in words]

        words = list(np.unique(words))

        return words

    @property
    @lru_cache
    def dictionary(self) -> pd.DataFrame:
        logging.log(logging.INFO, "Getting dictionary from document")
        tables = self.docx_tables
        data = []
        for table in tables:
            n_cols = len(table.columns)

            for row in table.rows:
                cells = row.cells
                values = [cell.text for cell in cells]
                for k in range(n_cols - 1):
                    try:
                        title = values[k]
                        content = values[k + 1]
                        if title != content and title != "" and content != "":
                            data.append(
                                {
                                    "title": title,
                                    "content": content,
                                    "orientation": "row",
                                }
                            )
                    except IndexError:
                        pass

        for table in tables:
            n_rows = len(table.rows)

            for col in table.columns:
                cells = col.cells
                values = [cell.text for cell in cells]
                for k in range(n_rows - 1):
                    try:
                        title = values[k]
                        content = values[k + 1]
                        if title != content and title != "" and content != "":
                            data.append(
                                {
                                    "title": title,
                                    "content": content,
                                    "orientation": "column",
                                }
                            )
                    except IndexError:
                        pass

        df = pd.DataFrame(data)
        df.drop_duplicates(inplace=True)

        return df

    @property
    @lru_cache
    def tables(self) -> List[pd.DataFrame]:
        logging.log(logging.INFO, "Getting tables from document")
        tables = self.docx_tables

        # Getting subset of tables that has a merged header
        splited_tables = []
        aux_table = []
        for k, table in enumerate(tables):
            for row in table.rows:
                values = [x.text for x in row.cells]
                if len(set(values)) == 1:
                    if len(values[0]) < 200:
                        # this is a merged table header with limited text lenght
                        splited_tables.append(aux_table)
                        aux_table = []
                else:
                    aux_table.append(values)

        dfs = []

        # Gettind data from conventional tables
        for table in tables:
            data = []
            headers = []
            for i, row in enumerate(table.rows):
                row = [cell.text for cell in row.cells]

                if i == 0:
                    headers = row
                else:
                    data.append(row)

            # TODO sometimes the lenfth of headers is different from the lenght of data
            # Check these cases and find a way to handle this situation
            df = pd.DataFrame(columns=headers, data=data)
            if len(df) > 0:
                dfs.append(df)

        # Getting data from splited tables
        for table in splited_tables:
            if len(table) == 0:
                continue

            header = table[0]
            data = table[1:]
            header_size = len(header)

            aux = []

            # Removing data with different length from header
            for line in data:
                if len(line) == header_size:
                    aux.append(line)
                else:
                    break

            df = pd.DataFrame(columns=header, data=aux)
            dfs.append(df)

        dfs = self.__merge_dfs(dfs)

        return dfs

    @property
    @lru_cache
    def docx_tables(self) -> List[Table]:
        """List of tables in docx document"""
        document = self.document
        tables = document.tables

        new_tables = tables
        count = 0

        # getting all inner tables
        while len(new_tables) > 0:
            count += 1
            aux = []
            for table in new_tables:
                for row in table.rows:
                    for cell in row.cells:
                        aux.extend(cell.tables)

            tables.extend(aux)
            new_tables = aux

            if count > 10:
                break

        return tables

    @property
    @lru_cache
    def document(self) -> Document:
        """Open document and creates a docx file if necessary

        Returns:
            - Document: word document object
        """
        folder, file_name = os.path.split(self.__file_path)
        self.file_name = self.__file_path.stem
        file_extension = self.__file_path.suffix[1:]

        if file_extension == "doc":
            docx_file_name = f"x_{file_name}x"
            docx_file_path = os.path.join(AUX_FOLDER_PATH, docx_file_name)

            if os.path.exists(docx_file_path):
                pass
            else:
                # Create a docx file if it yet does not exist in folder
                logging.info("Creating the docx file in the aux folder")

                # NOTE
                # If errors are found, do this
                # clear contents of C:\Users\<username>\AppData\Local\Temp\gen_py
                # C:\Users\u122004\AppData\Local\Temp\gen_py

                # Opening MS Word
                word = win32.gencache.EnsureDispatch("Word.Application")

                # NOTE this function does not accept the path as an object
                doc = word.Documents.Open(str(self.__file_path))
                doc.Activate()

                # Save and Close
                word.ActiveDocument.SaveAs(
                    docx_file_path, FileFormat=win32.constants.wdFormatXMLDocument
                )
                doc.Close(False)

            self.__file_path = docx_file_path

        document = Document(self.__file_path)
        return document

    def __merge_dfs(self, dfs: List[pd.DataFrame]) -> List[pd.DataFrame]:
        """Merge dataframes that has the same header and drop duplicated lines

        Args:
            - dfs (List[pd.DataFrame]): list of dataframes from doc extraction

        Returns:
            - List[pd.DataFrame]: list of merged dataframes
        """
        headers = ["&".join(df.columns.tolist()) for df in dfs]
        header_df = pd.DataFrame(headers, columns=["headers"])
        header_groups = header_df.groupby(by="headers")

        merged_dfs = []
        for _, df in header_groups:
            index = df.index.tolist()

            group = [dfs[i] for i in index]
            merged = pd.concat(group)
            merged.drop_duplicates(inplace=True)
            merged.reset_index(drop=True, inplace=True)
            merged_dfs.append(merged)

        return merged_dfs
