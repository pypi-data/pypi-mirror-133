#!python
#-*-- coding: utf-8 -*-

########################################################
#
# GTF2seq: a software for extracting sequences from
#     genome based on GTF file with variations
#
########################################################

import sys
__version__ = "0.1.0"


# general class containing info about a transcribed region.  Can come from UCSC Knowngenes (BED) or Ensembl GTF files currently 
# Originally written by Stephen N. Floor in Fall 2014 
# Modified on April 2, 2021
class Transcript: 
    def __init__(self):
        #properties defined in UCSC knowngenes 
        self.name = ''
        self.symbol = '' 
        self.chrom = ''
        self.strand = '' 
        self.txStart = 0
        self.txEnd = 0 
        self.cdsStart = 0
        self.cdsEnd = 0
        self.exonCt = 0
        self.exonStarts = []
        self.exonEnds = []
        self.exonLengths = []
        
        #meta properties to be computed during construction.  these are lists of BED first four field tuples with the exception of Len terms which are the length of the total region for the gene 
        self.utr5 = []
        self.utr5Len = 0
        self.utr5start = []
        self.utr5startLen = 0
        self.cds = []
        self.cdsLen = 0 
        self.utr3 = []
        self.utr3Len = 0
        self.exons = []
        self.exonsLen = 0
        self.introns = []
        self.intronsLen = 0
        self.promoter = []  ## promoter 1kb
        self.p15k = []  ### 1-5k upstream of promoter
        self.d1k = []  ## downstream 1k  
        self.coding = False

        ## For extracting the sequences
        self.utr5_tx = []
        self.utr5Len_tx = 0
        self.cds_tx = []
        self.cdsLen_tx = 0
        self.utr3_tx = []
        self.utr3Len_tx = 0
        self.exons_tx = []
        self.exonsLen_tx = 0
        self.introns_tx = []
        self.intronsLen_tx = 0
        self.promoter_tx = []  ## promoter 1kb
        self.p15k_tx = []  ### 1-5k upstream of promoter
        self.d1k_tx = []  ## downstream 1k


    def __str__(self):  #currently roughly knownGenes format with a second line containing metadata 
        return "%s\t%s\t%s\t%d\t%d\t%d\t%d\t%d\t%s\t%s\n%s\t%d\t%s\t%d\t%s\t%d\t%s\t%d\t%s\t%d\t%s" % (self.name, self.chrom, self.strand, self.txStart, self.txEnd, self.cdsStart, self.cdsEnd, self.exonCt, self.exonStarts, self.exonEnds, self.utr5, self.utr5Len, self.cds, self.cdsLen, self.utr3, self.utr3Len, self.exons, self.exonsLen, self.introns, self.intronsLen, self.coding)
    
#BED format output is goal.  Fields are optional after featureEnd 
# chrom    featureStart   featureEnd   nameOfLine   score(0-1000)   strand   thickStart  thickEnd  itemRGBtuple  blockCount  blockSizes   blockStarts 

    #this function returns a list of BED-formatted strings for the feature passed as region with multiple entries per region possible, one for each primitive (exon/intron) 
    #def bedFormat(self, region="exons", tx_id = "mm10"):
        #if (not self.coding and (region == "5utr" or region == "cds" or region == "3utr")):
        #    print("Transcript.py bedFormat error: noncoding transcripts do not have 5utr/cds/3utr")
        #    return []

        #returnVal = []

        #if (region == "5utr"):
        #    for chunk in self.utr5:
        #        #returnVal.append("%s\t%d\t%d\t%s\t%s\t%s" % (chunk[0], chunk[1], chunk[2], self.name, self.symbol, tx_id))

        #elif (region == "5utr_start"):
        #    for chunk in self.utr5start:
        #        #returnVal.append("%s\t%d\t%d\t%s\t%s\t%s" % (chunk[0], chunk[1], chunk[2], self.name, self.symbol, tx_id))

        #elif (region == "cds"):
        #    for chunk in self.cds:
        #        #returnVal.append("%s\t%d\t%d\t%s\t%s\t%s" % (chunk[0], chunk[1], chunk[2], self.name, self.symbol, tx_id))

        #elif (region == "3utr"):
        #    for chunk in self.utr3:
        #        #returnVal.append("%s\t%d\t%d\t%s\t%s\t%s" % (chunk[0], chunk[1], chunk[2], self.name, self.symbol, tx_id))

        #elif (region == "exons"):
        #    for chunk in self.exons:
        #        #returnVal.append("%s\t%d\t%d\t%s\t%s\t%s" % (chunk[0], chunk[1], chunk[2], self.name, self.symbol, tx_id))

        #elif (region == "introns"):
        #    for chunk in self.introns:
        #        #returnVal.append("%s\t%d\t%d\t%s\t%s\t%s" % (chunk[0], chunk[1], chunk[2], self.name, self.symbol, tx_id))
        #elif (region == "promoter"):
        #    #print("%s\t%d\t%d\t%s\t%s\t%s" % (self.chrom, self.promotStart, self.promotEnd, self.name, self.symbol, tx_id))
        #    returnVal.append("%s\t%d\t%d\t%s\t%s\t%s" % (self.chrom, self.promotStart, self.promotEnd, self.name, self.symbol, tx_id))
        #elif (region == "p1-5k"):
        #    returnVal.append("%s\t%d\t%d\t%s\t%s\t%s" % (self.chrom, self.p15kStart, self.p15kEnd, self.name, self.symbol, tx_id))
        #elif (region == "d1k"):
        #    returnVal.append("%s\t%d\t%d\t%s\t%s\t%s" % (self.chrom, self.d1kStart, self.d1kEnd, self.name, self.symbol, tx_id))

        #else:
        #    print("Transcript.py bedFormat error: currently only regions 5utr/cds/3utr/exons/introns are supported")
        #    

        #return returnVal

    def computeMetadata(self): 
    # -- begin computing metadata -- 

    # -- note: chose clarity of code and conditionals here over most efficient computation (i.e. some clauses may be redundant)

        if (self.strand == "+"): 
            ### Promoter 
            self.promotStart = self.txStart -1000 if self.txStart -1000 > 0 else 0
            self.promotEnd = self.txStart -1 if self.txStart -1 > 0 else 0
            ### 1-5K 
            self.p15kStart = self.txStart -5000 if  self.txStart -5000 > 0 else 0
            self.p15kEnd = self.txStart -1000 if  self.txStart -1000 >0 else 0
            ### d1k
            self.d1kStart = self.txEnd + 1
            self.d1kEnd = self.txEnd + 1000
        #print ("DBUG - exonCt %d i %d exonEnds[i] %d cdsStart %d exonStarts[i] %d cdsEnd %d") % \
            #    (self.exonCt, i, self.exonEnds[i], self.cdsStart, self.exonStarts[i], self.cdsEnd)
            for i in range (self.exonCt): 
                if (self.cdsStart != self.cdsEnd): # if this is a coding transcript
                    self.coding = True
                # -- first compute 5'utr, CDS, 3'utr regions --
                #case 1 - exon spans 5' UTR/CDS/3' UTR
                    if (self.exonStarts[i] < self.cdsStart and self.exonEnds[i] > self.cdsEnd):
                        self.utr5.append((self.chrom, self.exonStarts[i], self.cdsStart, self.name))
                        self.utr5Len += self.cdsStart - self.exonStarts[i]
                        self.utr5start.append((self.chrom, self.exonStarts[i], self.cdsStart, self.name)) # for now just append the 5' utr exons to the utr5start 
                        self.utr5startLen += self.cdsStart - self.exonStarts[i]
                        self.cds.append((self.chrom, self.cdsStart, self.cdsEnd, self.name))
                        self.cdsLen += self.cdsEnd - self.cdsStart
                        self.utr3.append((self.chrom, self.cdsEnd, self.exonEnds[i], self.name))
                        self.utr3Len += self.exonEnds[i] - self.cdsEnd
                #case 2 - exon spans 5' UTR/CDS junction
                    elif (self.exonStarts[i] < self.cdsStart and self.exonEnds[i] >= self.cdsStart):
                        self.utr5.append((self.chrom, self.exonStarts[i], self.cdsStart, self.name))
                        self.utr5Len += self.cdsStart - self.exonStarts[i]
                        self.utr5start.append((self.chrom, self.exonStarts[i], self.cdsStart, self.name)) 
                        self.utr5startLen += self.cdsStart  - self.exonStarts[i]
                        self.cds.append((self.chrom, self.cdsStart, self.exonEnds[i], self.name))
                        self.cdsLen += self.exonEnds[i]- self.cdsStart
                #case 3 - exon spans CDS/3'UTR junction 
                    elif (self.exonStarts[i] >= self.cdsStart and self.exonStarts[i] <= self.cdsEnd and self.exonEnds[i] > self.cdsEnd):
                        self.cds.append((self.chrom, self.exonStarts[i], self.cdsEnd, self.name))
                        self.cdsLen += self.cdsEnd - self.exonStarts[i]
                        self.utr3.append((self.chrom, self.cdsEnd, self.exonEnds[i], self.name))
                        self.utr3Len += self.exonEnds[i] - self.cdsEnd
                #case 4 - exon is 5' UTR only 
                    elif (self.exonStarts[i] < self.cdsStart and self.exonEnds[i] < self.cdsStart): 
                        self.utr5.append((self.chrom, self.exonStarts[i], self.exonEnds[i], self.name))
                        self.utr5Len += self.exonEnds[i] - self.exonStarts[i]
                        self.utr5start.append((self.chrom, self.exonStarts[i], self.exonEnds[i], self.name)) 
                        self.utr5startLen += self.exonEnds[i] - self.exonStarts[i]
                #case 5 - exon is CDS only
                    elif (self.exonStarts[i] >= self.cdsStart and self.exonEnds[i] <= self.cdsEnd):
                        self.cds.append((self.chrom, self.exonStarts[i], self.exonEnds[i], self.name))
                        self.cdsLen += self.exonEnds[i] - self.exonStarts[i]
                #case 6 - exon is 3' UTR only 
                    elif (self.exonStarts[i] > self.cdsEnd and self.exonEnds[i] > self.cdsEnd):
                        self.utr3.append((self.chrom, self.exonStarts[i], self.exonEnds[i], self.name))
                        self.utr3Len += self.exonEnds[i] - self.exonStarts[i]
                    else: 
                        print("Thar be dragons - Transcript computeMetadata + stranded gene region parsing")


            # -- generate combined exonic and intronic regions -- 
            #exons are easy 
                self.exons.append((self.chrom, self.exonStarts[i], self.exonEnds[i], self.name))
                self.exonsLen += self.exonEnds[i] - self.exonStarts[i]
            
            #print "DBUG2: i %d self.exonCt-1 %d self.exonEnds %s self.exonStarts %s" % (i, self.exonCt-1, self.exonEnds, self.exonStarts)
        
                if (i < self.exonCt - 1): # only compute introns for nonterminal exons
                # an intron is the region between the end of the current exon and start of the next 
                    self.introns.append((self.chrom, self.exonEnds[i], self.exonStarts[i+1], self.name))
                    self.intronsLen += self.exonStarts[i+1] - self.exonEnds[i] 


        elif (self.strand == "-"):
            ### Promoter
            self.promotStart = self.txEnd
            self.promotEnd = self.txEnd + 1000
            ### 1-5K
            self.p15kStart = self.txEnd + 1000
            self.p15kEnd = self.txEnd + 5000
            ### downstread 1k
            self.d1kStart = self.txStart - 1000 if  self.txStart - 1000 > 0 else 0
            self.d1kEnd = self.txStart - 1 if  self.txStart - 1 >0 else 0
     #uc001ach.2	    chr1    -	    910578  917473  911551  916546  5	    910578,911878,914260,916516,917444,	    911649,912004,916037,916553,917473,	    Q5SV97  uc001ach.2
            #	name		chrom	strand	txStart txEnd	cdsStart self.cdsEnd exonCt	exonStarts		exonEnds		proteinID  alignID 
            # for the minus strand everything is the same except the order of encountering regions is reversed
            # i.e. 3' UTR -> CDS -> 5' UTR 
            
            for i in range (self.exonCt): 
            #print ("DBUG - exonCt %d i %d self.exonEnds[i] %d self.cdsStart %d exonStarts[i] %d self.cdsEnd %d") % \
                #    (self.exonCt, i, self.exonEnds[i], self.cdsStart, self.exonStarts[i], self.cdsEnd)
                
                if (self.cdsStart != self.cdsEnd):
                    self.coding = True
                    #if self.symbol == "Mrpl15":
                    #    print self.cdsStart + "\t" + self.cdsEnd
                # -- first compute 5'utr, CDS, 3'utr regions --
                # -- this is the same as for + sense except 5' UTR and 3' UTR are swapped throughout
                #case 1 - exon spans 3' UTR/CDS/5' UTR
                    if (self.exonStarts[i] < self.cdsStart and self.exonEnds[i] > self.cdsEnd):
                        self.utr3.append((self.chrom, self.exonStarts[i], self.cdsStart, self.name))
                        self.utr3Len += self.cdsStart - self.exonStarts[i]
                        self.cds.append((self.chrom, self.cdsStart, self.cdsEnd, self.name))
                        self.cdsLen += self.cdsEnd - self.cdsStart
                        self.utr5.append((self.chrom, self.cdsEnd, self.exonEnds[i], self.name))
                        self.utr5Len += self.exonEnds[i] - self.cdsEnd
                        self.utr5start.append((self.chrom, self.cdsEnd, self.exonEnds[i], self.name))
                        self.utr5startLen += self.exonEnds[i] - (self.cdsEnd)
                #case 2 - exon spans 3' UTR/CDS junction
                    elif (self.exonStarts[i] < self.cdsStart and self.exonEnds[i] >= self.cdsStart):
                        self.utr3.append((self.chrom, self.exonStarts[i], self.cdsStart, self.name))
                        self.utr3Len += self.cdsStart - self.exonStarts[i]
                        self.cds.append((self.chrom, self.cdsStart, self.exonEnds[i], self.name))
                        self.cdsLen += self.exonEnds[i]- self.cdsStart
                #case 3 - exon spans CDS/5'UTR junction 
                    elif (self.exonStarts[i] >= self.cdsStart and self.exonStarts[i] <= self.cdsEnd and self.exonEnds[i] > self.cdsEnd):
                        self.cds.append((self.chrom, self.exonStarts[i], self.cdsEnd, self.name))
                        self.cdsLen += self.cdsEnd - self.exonStarts[i]
                        self.utr5.append((self.chrom, self.cdsEnd, self.exonEnds[i], self.name))
                        self.utr5Len += self.exonEnds[i] - self.cdsEnd
                        self.utr5start.append((self.chrom, self.cdsEnd, self.exonEnds[i], self.name))
                        self.utr5startLen += self.exonEnds[i] - (self.cdsEnd)
                #case 4 - exon is 3' UTR only 
                    elif (self.exonStarts[i] < self.cdsStart and self.exonEnds[i] < self.cdsStart): 
                        self.utr3.append((self.chrom, self.exonStarts[i], self.exonEnds[i], self.name))
                        self.utr3Len += self.exonEnds[i] - self.exonStarts[i]
                #case 5 - exon is CDS only
                    elif (self.exonStarts[i] >= self.cdsStart and self.exonEnds[i] <= self.cdsEnd):
                        self.cds.append((self.chrom, self.exonStarts[i], self.exonEnds[i], self.name))
                        self.cdsLen += self.exonEnds[i] - self.exonStarts[i]
                #case 6 - exon is 5' UTR only 
                    elif (self.exonStarts[i] > self.cdsEnd and self.exonEnds[i] > self.cdsEnd):
                        self.utr5.append((self.chrom, self.exonStarts[i], self.exonEnds[i], self.name))
                        self.utr5Len += self.exonEnds[i] - self.exonStarts[i]
                        self.utr5start.append((self.chrom, self.exonStarts[i] , self.exonEnds[i], self.name))
                        self.utr5startLen += self.exonEnds[i] - self.exonStarts[i]
                    else: 
                        print("Thar be dragons - Transcript computeMetadata - stranded gene region parsing")
                    
            #else: 
            #    print "- strand noncoding transcript"
                

            # -- generate combined exonic and intronic regions -- 
            #exons are easy 
                self.exons.append((self.chrom, self.exonStarts[i], self.exonEnds[i], self.name))
                self.exonsLen += self.exonEnds[i] - self.exonStarts[i]
            
                if (i < self.exonCt - 1): # only compute introns for nonterminal exons
                # an intron is the region between the end of the current exon and start of the next 
                    self.introns.append((self.chrom, self.exonEnds[i], self.exonStarts[i+1], self.name))
                    self.intronsLen += self.exonStarts[i+1] - self.exonEnds[i] 
                
        else:
            print("Thar be dragons - Transcript computeMetadata strand does not match + or -")
        

# input to createGTFTranscript below must be a list of dictionaries for each line of the input GTF file 
# these are created inside knowngenes_to_transcript_regions.py 

# example input: 

#[{'gene_name': 'DDX11L1', 'seqname': '1', 'end': '12227', 'start': '11869', 'frame': None, 'transcript_source': 'havana', 'feature': 'exon', 'exon_number': '1', 'exon_id': 'ENSE00002234944', 'tss_id': 'TSS15145', 'source': 'processed_transcript', 'gene_source': 'ensembl_havana', 'score': None, 'gene_biotype': 'pseudogene', 'gene_id': 'ENSG00000223972', 'transcript_id': 'ENST00000456328', 'transcript_name': 'DDX11L1-002', 'strand': '+'}, {'seqname': '1', 'end': '14409', 'start': '11869', 'frame': None, 'transcript_source': 'havana', 'feature': 'transcript', 'gene_id': 'ENSG00000223972', 'tss_id': 'TSS15145', 'source': 'processed_transcript', 'gene_source': 'ensembl_havana', 'score': None, 'gene_biotype': 'pseudogene', 'gene_name': 'DDX11L1', 'transcript_id': 'ENST00000456328', 'transcript_name': 'DDX11L1-002', 'strand': '+'}]

# keys for each dict:
#  gene_name
#  seqname
#  start
#  end
#  frame
#  transcript_source
#  feature
#  exon_number
#  exon_id
#  tss_id
#  source
#  gene_source
#  score
#  gene_biotype
#  gene_id
#  transcript_id
#  transcript_name
#  strand

def createGTFTranscript(gtfLines):
    """
    input: txDict[key] txDict for a specific transcript
    return: Transcript object
    """
    foo = Transcript()
    
    # these properties (better be) all identical for each entry in the list of dicts 
    
    first = gtfLines[0] 

    foo.name = first["transcript_id"]
    foo.symbol = first["gene_name"] if "gene_name" in first else first["transcript_id"]
    foo.chrom = first["seqname"]
    foo.strand = first["strand"]

    # now process all lines for this transcript ID 

    for dict in gtfLines: 
        
        # ensembl GTFs have special lines where feature = "transcript" and feature = "CDS" that define the transcript and CDS start/ends, respectively 
    
        # GTF files are closed intervals while BED are right-open-left-closed, so --- 
        #   need to subtract one from all start coordinates? seems counterintuitive maybe the input genome.fa is zero based? 

        if (dict["feature"] == "exon"):
            ### Get the start and end position of transcript
            if (foo.txStart == 0 or int(dict["start"]) < foo.txStart):
                foo.txStart = int(dict["start"]) - 1
            if (foo.txEnd == 0 or int(dict["end"]) > foo.txEnd):
                foo.txEnd = int(dict["end"])
            
            ### 
            foo.exonCt += 1 
            foo.exonStarts.append(int(dict["start"]) - 1)
            foo.exonEnds.append(int(dict["end"]))


        if (dict["feature"] == "CDS"):
            #print dict
            if (foo.cdsStart == 0 or int(dict["start"]) < foo.cdsStart):
                foo.cdsStart = int(dict["start"]) - 1
            if (foo.cdsEnd== 0 or int(dict["end"]) > foo.cdsEnd):
                foo.cdsEnd = int(dict["end"])
            #if foo.symbol == "Mrpl15":
            #    print str(foo.cdsStart) +"\t"+str(foo.cdsEnd) 
            
        ##if (dict["feature"] == "exon"):  comment out by 
        #    foo.exonCt += 1
        #
        #    foo.exonStarts.append(int(dict["start"]) - 1)
        #    foo.exonEnds.append(int(dict["end"]))

    foo.exonStarts = sorted(foo.exonStarts)
    foo.exonEnds = sorted(foo.exonEnds) 

    foo.computeMetadata() 
    return foo 
#####################################################################################

import argparse
import logging 
from textwrap import dedent
#from Transcript import *
import re 
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

from pyfaidx import Fasta
## pyfaidx 

import vcf
import logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s :: %(levelname)s :: %(message)s')


###########################GTF############################################
#VCFREADER = None
SAMPLE = None
SNPRECORD = {}

from collections import defaultdict
import gzip
import re

### Colors for different features
color_p15k = RGBColor(230, 75, 53)  # "Cinnabar" = "#E64B35"
color_promot = RGBColor(126, 97, 72)   # "RomanCoffee" = "#7E6148"  
color_utr5 = RGBColor(77, 187, 213) # "Shakespeare" = "#4DBBD5",
color_exon = RGBColor(0, 160, 135)  # "PersianGreen" = "#00A087"
color_cds = RGBColor(60, 84, 136)   # "Chambray" = "#3C5488",
color_intron = RGBColor(243, 155, 127)   #"Apricot" = "#F39B7F"
color_utr3 = RGBColor(132, 147, 180) #"WildBlueYonder" = "#8491B4",
color_d1k = RGBColor(145, 209, 191) #"MonteCarlo" = "#91D1C2", "Monza" = "#DC0000",
# "Sandrift" = "#B09C85"

color_dict = {}
color_dict["Upstream 1-5k"] = RGBColor(230, 75, 53) 
color_dict["Promoter (upstream 1k)"] = RGBColor(126, 97, 72) 
color_dict["5' UTR"] = RGBColor(77, 187, 213)
color_dict["Exon"] = RGBColor(0, 160, 135)
color_dict["CDS"] = RGBColor(60, 84, 136)
color_dict["Intron"] = RGBColor(243, 155, 127)
color_dict["3' UTR"] = RGBColor(132, 147, 180) 
color_dict["Downstream 1-5k"] = RGBColor(145, 209, 191)

GTF_HEADER  = ['seqname', 'source', 'feature', 'start', 'end', 'score', 'strand', 'frame']
R_SEMICOLON = re.compile(r'\s*;\s*')
R_COMMA     = re.compile(r'\s*,\s*')
R_KEYVALUE  = re.compile(r'(\s+|\s*=\s*)')

def gtf_reader(filename):
    """
    Open an optionally gzipped GTF file and generate a dict for each line.
    This line is iterable. 
    """
    fn_open = gzip.open if filename.endswith('.gz') else open

    with fn_open(filename) as fh:
        for line in fh:
            if line.startswith('#'):
                continue
            else:
                yield _parse_gtf_line(line)


def _parse_gtf_line(line):
    """Parse a single GTF line and return a dict.
    """
    result = {}
    fields = line.rstrip().split('\t')
    for i, col in enumerate(GTF_HEADER):  ## i index col include the content 
        #logging.debug(i)
        #logging.debug(col)
        result[col] = _get_value(fields[i])
    # INFO field consists of "key1=value;key2=value;...".
    infos = re.split(R_SEMICOLON, fields[8])    ## split last feilds 
    for i, info in enumerate(infos, 1):
        # It should be key="value".
        try:
            key, _, value = re.split(R_KEYVALUE, info)
        # But sometimes it is just "value".
        except ValueError:
            key = 'INFO{}'.format(i)
            value = info
        # Ignore the field if there is no value.
        if value:
            result[key] = _get_value(value)

    return result

def _get_value(value):
    if not value:
        return None
    # Strip double and single quotes.
    value = value.strip('"\'')
    # Return a list if the value has a comma.
    if ',' in value:
        value = re.split(R_COMMA, value)
    # These values are equivalent to None.
    elif value in ['', '.', 'NA']:
        return None

    return value
#########################END GTF##################################################

def reverse_complement(seq):
    old_chars = "ACGT"
    replace_chars = "TGCA"
    tab = str.maketrans(old_chars,replace_chars) 
    return seq.translate(tab)[::-1] 

## process the transcript ID list file
def read_transcriptID(transcriptid_list):
    transcriptID_list = {}
    with open(transcriptid_list) as ID:
        for line in ID:
            line = line.rstrip("\n")
            line = re.sub("\s+", "", line)
            transcriptID_list[line] = 0
    return transcriptID_list 

def paragraph_seq_add_run(paragraph_seq, base, coor, color):
    run_seq = paragraph_seq.add_run(base)
    font_seq = run_seq.font
    font_seq.color.rgb = color
    ### If its a 
    if coor in SNPRECORD:
        #print(coor)
        font_seq.highlight_color = WD_COLOR_INDEX.GRAY_25 ## Lighgray
        if SNPRECORD[coor]["IMPACT"] == "HIGH":
            font_seq.bold = True
        if SNPRECORD[coor]["IMPACT"] == "MODERATE":
            font_seq.italic = True
        if SNPRECORD[coor]["MISS"] == True:
            ## https://python-docx.readthedocs.io/en/latest/dev/analysis/features/text/font-highlight-color.html
            font_seq.underline = True

def get_element_sequence(genome_seq, tx, range_start, range_end, step, paragraph_seq, color):
    """
    Extract sequences from genome based on coordinates. 
    Parameters:
        genome_seq: Fasta object; 
        tx: GTFTranscript object
        range_start: start position
        range_end: end position
        step: strand information: 1 for + and -1 for -  
        paragraph_seq: word paragraph object
        color: which color to color the word text 
    """
    #print([genome_seq, tx, range_start, range_end, step, paragraph_seq, color]) 
    seq = ""
    if step ==1:    ## 
        for i in range(range_start, range_end + 1, step):
            base = genome_seq[tx.chrom][i-1:i].seq
            coor = "\t".join([tx.chrom, str(i)])
            if coor in SNPRECORD:
                base = SNPRECORD[coor]["GENOTYPE"]
            paragraph_seq_add_run(paragraph_seq, base, coor, color)
            seq = seq + base
    else:
        for i in range(range_start, range_end - 1, step):
            base = genome_seq[tx.chrom][i-1:i].seq
            coor = "\t".join([tx.chrom, str(i)])
            if coor in SNPRECORD:
                base = SNPRECORD[coor]["GENOTYPE"]
            base_rc = reverse_complement(base)
            paragraph_seq_add_run(paragraph_seq, base_rc, coor, color)
            seq = seq+base_rc
            #print([tx.chrom, base])
    #print(seq)
    #run_seq = paragraph_seq.add_run(seq)
    #font_seq = run_seq.font
    #font_seq.color.rgb = color
    return seq

def process_seq(genome_seq, tx, doc):
    #### Extract seqeunce 
    paragraph_id = doc.add_paragraph('>' + tx.name + " " + str(tx.coding))
    paragraph_seq = doc.add_paragraph('') 
    
    if tx.strand == "+": ## positive strand 
        seq_p15k = get_element_sequence(genome_seq, tx, tx.p15kStart, tx.p15kEnd, 1, paragraph_seq, color_p15k)
        seq_promot = get_element_sequence(genome_seq, tx, tx.promotStart, tx.promotEnd, 1, paragraph_seq, color_promot)
        count_intron = 0
        if tx.coding is True:  ## Protein coding genes
            count_intron_utr = len(tx.utr5) - 1 
            for i in range(0, len(tx.utr5), 1):  #[('1', 2170279, 2170492, 'Zm0'), ('1', 2170756, 2170983, 'Zm000")]
                seq_utr5 = get_element_sequence(genome_seq, tx, tx.utr5[i][1], tx.utr5[i][2], 1, paragraph_seq, color_utr5)
                count_intron_utr = count_intron_utr -1
                if count_intron_utr >= 0:  ## there are more than one utr5 sections,meaning there are introns in 5' UTR
                    if options.exclude_intron is not True: 
                        seq_intron = get_element_sequence(genome_seq, tx, tx.utr5[i][1], tx.introns[i][2], tx.utr5[i][1], 1, paragraph_see, color_intron)
                    count_intron = count_intron + 1
                else: ## last 5'UTR section
                    if tx.introns[i][1] == tx.utr5[i][2]:   ## ## there is a intron between 5' UTR and CDS
                        if options.exclude_intron is not True:
                            seq_intron = get_element_sequence(genome_seq, tx, tx.introns[i][1], tx.introns[i][2], 1, paragraph_seq, color_intron)
                        count_intron = count_intron + 1
            ### cds
            count_intron_cds = len(tx.cds) -1 
            for i in range(0, len(tx.cds)):
                seq_utr5 = get_element_sequence(genome_seq, tx, tx.cds[i][1], tx.cds[i][2], 1, paragraph_seq, color_cds)
                # there are more than 1 CDS and there are more than one remaining
                if count_intron_cds >=0:
                    if options.exclude_intron is not True:
                        seq_intron = get_element_sequence(genome_seq, tx, tx.introns[count_intron][1], tx.introns[count_intron][2], 1, paragraph_seq, color_intron)
                    count_intron = count_intron + 1
                else:
                    if tx.cds[i][2] == tx.introns[count_intron][1]:
                        if options.exclude_intron is not True:
                            seq_intron = get_element_sequence(genome_seq, tx, tx.introns[count_intron][1], tx.introns[count_intron][2], 1, paragraph_seq, color_intron)
                        count_intron = count_intron + 1
             ### 3' utr                       
            count_intron_utr3 = len(tx.utr3) - 1
            for i in range(0, len(tx.utr3), 1):  #[('1', 2170279, 2170492, 'Zm0'), ('1', 2170756, 2170983, 'Zm000")]
                seq_utr3 = get_element_sequence(genome_seq, tx, tx.utr3[i][1], tx.utr3[i][2], 1, paragraph_seq, color_utr3)
                count_intron_utr3 = count_intron_utr3 -1
                if count_intron_utr3 >= 0:  ## there are more than one utr5 sections,meaning there are introns in 5' UTR
                    if options.exclude_intron is not True:
                        seq_intron = get_element_sequence(genome_seq, tx, tx.introns[count_intron][1], tx.introns[count_intron][2], 1, paragraph_see, color_intron)
                    count_intron = count_intron + 1 
                        
        seq_d1k = get_element_sequence(genome_seq, tx, tx.d1kStart, tx.d1kEnd, 1, paragraph_seq, color_d1k)
    else:  ## reverse "-"
        seq_p15k = get_element_sequence(genome_seq, tx, tx.p15kEnd, tx.p15kStart, -1, paragraph_seq, color_p15k)
        seq_promot = get_element_sequence(genome_seq, tx, tx.promotEnd, tx.promotStart, -1, paragraph_seq, color_promot)
        count_intron = 0 
        if tx.coding is True:  ## Protein coding genes
            ### utr5 
            for i in range(len(tx.utr5) - 1, -1, -1):  #[('1', 2170279, 2170492, 'Zm0'), ('1', 2170756, 2170983, 'Zm000")]
                seq_utr5 = get_element_sequence(genome_seq, tx, tx.utr5[i][2], tx.utr5[i][1], -1, paragraph_seq, color_utr5)
                intron_index = i  
                if i > 0:  ## there are more than one utr5 sections,meaning there are introns in 5' UTR 
                    if options.exclude_intron is not True:
                        seq_intron = get_element_sequence(genome_seq, tx, tx.introns[i][2], tx.utr5[i][1], -1, paragraph_seq, color_intron)
                    count_intron = count_intron + 1
                else:      ## last5' UTR section 
                    if  tx.introns[i][2] == tx.utr5[i][1]:  ## there is a intron between 5' UTR and CDS
                        if options.exclude_intron is not True:
                            seq_intron = get_element_sequence(genome_seq, tx, tx.introns[i][2], tx.intron[i][1], -1, paragraph_seq, color_intron)
                        count_intron = count_intron + 1 
            ### cds
            #print("intron_num: ")
            #print(count_intron)
            for i in range(len(tx.cds) -1, -1, -1):
                #print(i)
                seq_utr5 = get_element_sequence(genome_seq, tx, tx.cds[i][2], tx.cds[i][1], -1, paragraph_seq, color_cds) 
                # there are more than 1 CDS and there are more than one remaining
                if i > 0:
                    if options.exclude_intron is not True:
                        seq_intron = get_element_sequence(genome_seq, tx, tx.introns[count_intron][2], tx.introns[count_intron][1], -1, paragraph_seq, color_intron)
                    count_intron = count_intron + 1 
                else:  
                    #print(count_intron) 
                    #print(i)
                    if count_intron <= len(tx.introns)-1 and tx.introns[count_intron][2] == tx.cds[i][1]:
                        if options.exclude_intron is not True:
                            seq_intron = get_element_sequence(genome_seq, tx, tx.introns[count_intron][2], tx.introns[count_intron][1], -1, paragraph_seq, color_intron)
                        count_intron = count_intron + 1
            ## 3' utr
            for i in range(len(tx.utr3) - 1, -1, -1):  #[('1', 2170279, 2170492, 'Zm0'), ('1', 2170756, 2170983, 'Zm000")]
                seq_utr3 = get_element_sequence(genome_seq, tx, tx.utr3[i][2], tx.utr3[i][1], -1, paragraph_seq, color_utr3)
                intron_index = i
                if i > 0:  ## there are more than one utr5 sections,meaning there are introns in 5' UTR
                    if options.exclude_intron is not True:
                        seq_intron = get_element_sequence(genome_seq, tx, tx.introns[count_intron][2], tx.introns[count_intron][1], -1, paragraph_seq, color_intron)
                    count_intron = count_intron + 1
        else:   ## There is no UTR    
            #logging.info("XX") 
            ## For each exon 
            for i in range(len(tx.exons)-1, -1, -1):  #[('1', 2170279, 2170492, 'Zm0'), ('1', 2170756, 2170983, 'Zm000")]
                logging.info(i)
                seq_exon  = get_element_sequence(genome_seq, tx, tx.exons[i][2], tx.exons[i][1], -1, paragraph_seq, color_exon)
                logging.info(seq_exon)
                intron_index = len(tx.introns) - count_intron - 1
                if intron_index >= 0:
                    if options.exclude_intron is not True:
                        seq_intron = get_element_sequence(genome_seq, tx, tx.introns[intron_index][2], tx.exons[intron_index][1], -1, paragraph_seq, color_intron)
                        #logging.info("Intron:" + str(intron_index) + seq_intron)
                    count_intron = count_intron + 1 
        seq_d1k = get_element_sequence(genome_seq, tx, tx.d1kEnd, tx.d1kStart, -1, paragraph_seq, color_d1k)

def output_legend(doc, options):
    p = doc.add_paragraph('Legend information')
    for key in color_dict:
        p = doc.add_paragraph('')
        run_seq = p.add_run(key)
        font_seq = run_seq.font
        font_seq.color.rgb = color_dict[key]
    if options.vcf is not None:
        p = doc.add_paragraph('')
        run_seq = p.add_run('SNP position is highlighted in lightgray')
        font_seq = run_seq.font
        font_seq.highlight_color = WD_COLOR_INDEX.GRAY_25 ## Lighgray

        #######################################
        p = doc.add_paragraph('')
        run_seq = p.add_run('HIGH impact SNPs is bolded')
        font_seq = run_seq.font
        font_seq.highlight_color = WD_COLOR_INDEX.GRAY_25 ## Lighgray
        font_seq.bold = True

        ########################################
        p = doc.add_paragraph('')
        run_seq = p.add_run('Moderate impact SNPs is italic')
        font_seq = run_seq.font
        font_seq.highlight_color = WD_COLOR_INDEX.GRAY_25 ## Lighgray
        font_seq.italic = True
        ########################################
        if options.sample is not None:
            p = doc.add_paragraph('')
            run_seq = p.add_run('SNP position with missing genotype is shown with underline')
            font_seq = run_seq.font
            font_seq.underline = True

def print_check_tx(tx, key):
    print("================: \t" + key + "\t" + tx.strand)
    print("strand: ")
    print(tx.strand)
    print("p15k")
    print([tx.p15kStart, tx.p15kEnd])
    print("p_promoter")
    print([tx.promotStart, tx.promotEnd])
    print("cds:")
    print(tx.cds)
    print(tx.cdsLen)
    print("intron")
    print(tx.introns)
    print("exon")
    print(tx.exons)
    print("utr5")
    print(tx.utr5)
    print("utr3")
    print(tx.utr3)
    print("downstream 1k")    
    print([tx.d1kStart, tx.d1kEnd])

def test_get_snp_record(vcf_reader,tx):
    
    for record in vcf_reader.fetch('1', 1110695, 1230237): 
        coor = "\t".join([record.CHROM, str(record.POS)])
        SNPRECORD[coor] = [record.REF]
        print(record)
        print(record.REF)
        print(record.genotype("Zea_mays_Mo17").gt_bases)

        gt_base = record.genotype("Zea_mays_Mo17").gt_bases
        #gt_base = re.sub("|", "", gt_base)
        gt_base = re.sub("/", "", gt_base)
        #print(gt_base)

def get_snp_record(vcf_reader, tx):
    for record in vcf_reader.fetch(tx.chrom, tx.txStart, tx.txEnd):
        coor = "\t".join([record.CHROM, str(record.POS)])
        SNPRECORD[coor] = {}
        SNPRECORD[coor]["GENOTYPE"] = record.REF
        SNPRECORD[coor]["IMPACT"] = None
        SNPRECORD[coor]["MISS"] = False
        if SAMPLE is not None:
            gt_bases = record.genotype(SAMPLE).gt_bases 
            if gt_bases is None:
                SNPRECORD[coor]["MISS"] = True
            else:
                gt_base = re.sub("|", "", gt_base)
                gt_base = re.sub("/", "", gt_base)
                gt_base = re.sub(record.REF, "", gt_base)
                SNPRECORD[coor]["GENOTYPE"] = gt_base[0:1]
        #print(record.INFO)
        if "ANN" not in record.INFO:
            logging.warning("No snpEff anntation detected. Please make sure you are fine with this.")
        else:
            #print(record.INFO["ANN"])
            count = 0 
            for impact in record.INFO["ANN"]:
                #print(tx.name)
                #G|5_prime_UTR_variant|MODIFIER|Zm00001d027240|Zm00001d027240|transcript|Zm00001d027240_T004|protein_coding|4/5|c.-315T>C|||||315|
                ele = impact.split("|")
                #print(ele[2])
                #print(ele[6])
                if tx.name == ele[6]:
                    SNPRECORD[coor]["IMPACT"] = "MODERATE"  # ele[2] 
                    count = count + 1
            if count == 0:  
                logging.warning("No impacted detected for ." + tx.name + ". Plesae double check your VCF file")
        
def main(options):
    # the issue here is that lines for various transcripts may be interleaved,
    # so can either create lots of objects, or a giant dict. opted for giant dict.        
    document = Document() 
    vcf_reader = ""

    ### Check whether vcf file is provided or not
    if options.vcf is not None:
        logging.info("Reading SNP information from " + options.vcf)
        vcf_reader= vcf.Reader(filename=options.vcf)
        ## Check if the sample exists or not in the vcf file
        if options.sample is not None and options.sample not in vcf_reader.samples:
            sys.exit(options.sample + " not detected in vcf file: " + options.vcf)
        SAMPLE = options.sample 
    else:
        logging.info("SNP information is not provided. ")  
    
    #print(vcf_reader)
    transcriptID_list = read_transcriptID(options.transcriptid_list) ## Dictionary
    logging.info("Reading genome.")
    genome_seq = Fasta(options.fasta)
    record_detected_transcript = {} 
    txDict = defaultdict(list)
    genesRead = 0 
    for line in gtf_reader(options.gtf):
        # only want to read in lines corresponding to these features
        #print(line)
        if line["feature"] in ["exon", "CDS", "start_codon", "stop_codon"]:
            #print(line)
            txDict[line["transcript_id"]].append(line)
            genesRead += 1
            if (not genesRead % 25000):
                print("\tProcessed %d lines..." %  genesRead)
    logging.info("Dictionary for transcript built done.") 
    # now create a Transcript object for each transcript and output it 
    # transcrippt ID => Key
    for tID in transcriptID_list:
        logging.info(tID)
        if tID in txDict:
            tx = createGTFTranscript(txDict[tID])
            #print_check_tx(tx, tID)
            SNPRECORD = get_snp_record(vcf_reader, tx)
          
            process_seq(genome_seq, tx, document)
        else:
            logging.warning("'"+tID+"' " +"is not detected in the gtf file.")
    #for key in txDict: 
    #    logging.info(key)
    #    tx = createGTFTranscript(txDict[key])
    #    #Zm00001d027230\|Zm00001d027231
    #    #if "Zm00001d027230_T001" == key or "Zm00001d027231_" in key:
    #    #if "Zm00001d027240_T001" == key:
    #    if key in transcriptID_list:
    #        #process_seq(genome_seq, tx)
    #        record_detected_transcript[key] = 1
    #        print_check_tx(tx, key)
    #        process_seq(genome_seq, tx, document)
    output_legend(document, options)
    document.save(options.output)
    
if __name__ == '__main__':
    parser = argparse.ArgumentParser(formatter_class=argparse.RawTextHelpFormatter, description=dedent("""\
    """))    
    parser.add_argument('-g', '--gtf', type=str, help='Genome annotation file in GTF format', required=True)
    parser.add_argument('-f', '--fasta', type=str, help='Genome sequences in FASTA format', required=True)
    parser.add_argument('-t', '--transcriptid_list', type=str, help='List of transcript IDs', required=True)
    parser.add_argument('--vcf', type=str, help='VCF files with snpEff annotation. ', required=False)
    parser.add_argument('--sample', type=str, help='Sample/individual ID if you want to output genotypes in the sequence for a specific sample/individual.', required=False)
    parser.add_argument('--exclude_intron', action='store_true', help='Exclude intron sequences in the output') 
    parser.add_argument('-o', '--output', help='Output file in Word format', default="transcript_seq.docx")
    options = parser.parse_args()
    main(options)
    logging.info("Done!")
 
