import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

from pedurma.utils import translate_tib_number


def split_text(content):

    chunks = re.split(r"(\(\d+\) <.*?>)", content)

    return chunks


def create_docx_with_footnotes(text_id, collated_text, path):
    chunks = split_text(collated_text)
    document = Document()
    p = document.add_paragraph()

    for chunk in chunks:
        if chunk and "<" in chunk:
            chunk = re.sub(r"\(\d+\) ", "", chunk)
            super_text = p.add_run(chunk)
            super_text.font.superscript = True
            super_text.font.name = "Jomolhari"
        else:
            normal_text = p.add_run(chunk)
            normal_text.font.name = "Jomolhari"
    output_path = path / f"{text_id}_format_01.docx"
    document.save(str(output_path))
    return output_path


def get_pages(text):
    result = []
    pg_text = ""
    pages = re.split(r"(\d+-\d+)", text)
    for i, page in enumerate(pages[:-1]):
        if i % 2 == 0:
            pg_text += page
        else:
            pg_text += page
            result.append(pg_text)
            pg_text = ""
    return result


def parse_page(page, page_num):
    page_ann = re.search(r"(\d+-\d+)", page)[0]
    page = page.replace(page_ann, "")
    page_html = "<p>"
    chunks = split_text(page)
    for chunk in chunks:
        if chunk and "<" in chunk:
            footnote_number = re.search(r"\((\d+)\) <.*?>", chunk).group(1)
            footnote_number = translate_tib_number(footnote_number)
            page_html += f'<a href="#footnote-{page_num}-{footnote_number}"><sup>[{footnote_number}]</sup></a>'
        else:
            page_html += chunk
    page_html += f"</p><p>{page_ann}</p><p>------------------------</p>"
    notes = re.finditer(r"\((\d+)\) <(.*?)>", page)
    for note in notes:
        note_num = translate_tib_number(note.group(1))
        page_html += f'<p id="footnote-{page_num}-{note_num}"><sup>[{note_num}]</sup> {note.group(2)}</p>'
    page_html += '<div style="page-break-after: always"></div>'
    return page_html


def creat_docx_footnotes_at_end_of_page(text_id, collated_text, path):
    collated_text_html = ""
    pages = get_pages(collated_text)
    for page_num, page in enumerate(pages, 1):
        collated_text_html += parse_page(page, page_num)
    html_path = path / f"{text_id}.html"
    html_path.write_text(collated_text_html, encoding="utf-8")
    output_path = path / f"{text_id}.docx"
    os.system(f"ebook-convert {html_path} {output_path} --docx-no-toc")
    html_path.unlink()
    return output_path


def get_docx_text(text_id, preview_text, output_path=None, type_="with_footnotes"):
    if not output_path:
        (Path.home() / ".collation_docx").mkdir(parents=True, exist_ok=True)
        output_path = Path.home() / ".collation_docx"
    collated_text = ""
    for vol_id, text in preview_text.items():
        collated_text += f"{text}\n\n"
    collated_text = collated_text.replace("\n", "")
    collated_text = re.sub(r"(\d+-\d+)", r"\n\g<1>\n", collated_text)
    if type_ == "with_footnotes":
        docx_path = create_docx_with_footnotes(text_id, collated_text, output_path)
    else:
        docx_path = creat_docx_footnotes_at_end_of_page(
            text_id, collated_text, output_path
        )
    return docx_path
